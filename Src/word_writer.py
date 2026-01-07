import re
from pathlib import Path
from typing import List, Tuple

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt
except ImportError:
    Document = None
    qn = None
    Pt = None

HEADER_RE = re.compile(r"^([一二三四五六七八九十]+、|\d+[\.、]|[（(][一二三四五六七八九十0-9]+[)）])")
LABEL_LINE_RE = re.compile(r"^([^：\s]{1,8}：)(.*)$")
LABEL_INLINE_MAX = 20
PARAGRAPH_ENDINGS = ("。", "！", "？", "；", ":", "：", ".", "!", "?", ";")
MD_CODE_BLOCK_RE = re.compile(r"```.*?```", re.S)
MD_INLINE_CODE_RE = re.compile(r"`([^`]+)`")
MD_BOLD_RE = re.compile(r"(\*\*|__)(.+?)(\1)")
MD_ITALIC_RE = re.compile(r"(\*|_)([^*_]+?)(\1)")
MD_HEADING_RE = re.compile(r"^\s{0,3}#{1,6}\s+")
MD_HEADING_LINE_RE = re.compile(r"^\s{0,3}(#{1,6})\s+(.+)$")
MD_BLOCKQUOTE_RE = re.compile(r"^\s{0,3}>\s?")
MD_UL_RE = re.compile(r"^\s*[-*+]\s+")
MD_HR_RE = re.compile(r"^\s*([-*_]\s*){3,}$")
MD_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
MD_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
MD_HTML_TAG_RE = re.compile(r"<[^>]+>")
MD_ORDERED_LIST_RE = re.compile(r"^\s*(\d+)[\.\)、]\s+(.+)$")
MD_UNORDERED_LIST_RE = re.compile(r"^\s*[-*+]\s+(.+)$")
MD_TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*[:\-]+\s*(\|\s*[:\-]+\s*)+\|?\s*$")


def is_standalone_line(line: str) -> bool:
    if line.endswith(("：", ":")):
        return True
    if HEADER_RE.match(line):
        return True
    if line.startswith(("•", "-", "*")):
        return True
    return False


def ends_paragraph(line: str) -> bool:
    return line.endswith(PARAGRAPH_ENDINGS)


def join_lines(prev: str, next_line: str) -> str:
    if not prev:
        return next_line
    prev_tail = prev[-1]
    next_head = next_line[0]
    if prev_tail.isascii() and next_head.isascii() and prev_tail.isalnum() and next_head.isalnum():
        return f"{prev} {next_line}"
    return f"{prev}{next_line}"


def normalize_text_to_paragraphs(text: str) -> List[str]:
    paragraphs: List[str] = []
    current = ""
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            if current:
                paragraphs.append(current)
                current = ""
            continue
        label_match = LABEL_LINE_RE.match(line)
        if label_match:
            if current:
                paragraphs.append(current)
                current = ""
            rest = label_match.group(2).strip()
            if not rest:
                paragraphs.append(line)
                continue
            if ends_paragraph(line) or len(rest) <= LABEL_INLINE_MAX:
                paragraphs.append(line)
                continue
            current = line
            if ends_paragraph(current):
                paragraphs.append(current)
                current = ""
            continue
        if is_standalone_line(line):
            if current:
                paragraphs.append(current)
                current = ""
            paragraphs.append(line)
            continue
        if current:
            if ends_paragraph(current):
                paragraphs.append(current)
                current = line
            else:
                current = join_lines(current, line)
        else:
            current = line
        if ends_paragraph(current):
            paragraphs.append(current)
            current = ""
    if current:
        paragraphs.append(current)
    return paragraphs


def _clean_inline(text: str) -> str:
    cleaned = MD_INLINE_CODE_RE.sub(r"\1", text)
    cleaned = MD_IMAGE_RE.sub(r"\1 (\2)", cleaned)
    cleaned = MD_LINK_RE.sub(r"\1 (\2)", cleaned)
    cleaned = MD_ITALIC_RE.sub(r"\2", cleaned)
    cleaned = MD_HTML_TAG_RE.sub("", cleaned)
    return cleaned


def parse_inline_segments(text: str) -> List[Tuple[str, bool]]:
    segments: List[Tuple[str, bool]] = []
    if not text:
        return segments
    text = _clean_inline(text)
    cursor = 0
    for match in MD_BOLD_RE.finditer(text):
        if match.start() > cursor:
            segment = text[cursor:match.start()]
            if segment:
                segments.append((segment, False))
        bold_text = match.group(2)
        if bold_text:
            segments.append((bold_text, True))
        cursor = match.end()
    if cursor < len(text):
        segment = text[cursor:]
        if segment:
            segments.append((segment, False))
    return segments


def _apply_run_font(run) -> None:
    run.font.bold = run.bold
    run.font.name = "SimSun"
    if qn is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def add_runs_with_style(paragraph, text: str) -> None:
    for segment, is_bold in parse_inline_segments(text):
        run = paragraph.add_run(segment)
        run.bold = is_bold
        _apply_run_font(run)


def split_table_row(line: str) -> List[str]:
    row = line.strip().strip("|")
    cells = [cell.strip() for cell in row.split("|")]
    return [cell for cell in cells if cell != ""]


def is_heading_candidate(line: str) -> bool:
    if MD_ORDERED_LIST_RE.match(line) or MD_UNORDERED_LIST_RE.match(line):
        return False
    if "：" in line or ":" in line:
        return False
    if len(line.strip()) > 30:
        return False
    return HEADER_RE.match(line.strip()) is not None


def markdown_to_text(text: str) -> str:
    if not text:
        return ""
    stripped = MD_CODE_BLOCK_RE.sub(lambda m: m.group(0).strip("`"), text)
    stripped = MD_IMAGE_RE.sub(r"\1 (\2)", stripped)
    stripped = MD_LINK_RE.sub(r"\1 (\2)", stripped)
    stripped = MD_INLINE_CODE_RE.sub(r"\1", stripped)
    stripped = MD_BOLD_RE.sub(r"\2", stripped)
    stripped = MD_ITALIC_RE.sub(r"\2", stripped)

    lines = []
    for raw_line in stripped.splitlines():
        line = raw_line.strip()
        if not line:
            lines.append("")
            continue
        if MD_HR_RE.match(line):
            continue
        line = MD_HEADING_RE.sub("", line)
        line = MD_BLOCKQUOTE_RE.sub("", line)
        line = MD_UL_RE.sub("• ", line)
        line = MD_HTML_TAG_RE.sub("", line)
        lines.append(line)

    return "\n".join(lines).strip()


def parse_markdown_blocks(text: str) -> List[Tuple[str, object]]:
    lines = text.splitlines()
    blocks: List[Tuple[str, object]] = []
    i = 0
    while i < len(lines):
        raw_line = lines[i].rstrip("\n")
        line = raw_line.rstrip()
        if not line.strip():
            blocks.append(("blank", ""))
            i += 1
            continue
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            code_text = "\n".join(code_lines).strip()
            if code_text:
                blocks.append(("code", code_text))
            continue
        heading_match = MD_HEADING_LINE_RE.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            blocks.append(("heading", (level, heading_match.group(2).strip())))
            i += 1
            continue
        if "|" in line and i + 1 < len(lines) and MD_TABLE_SEPARATOR_RE.match(lines[i + 1]):
            header = split_table_row(line)
            i += 2
            rows = []
            while i < len(lines) and "|" in lines[i]:
                rows.append(split_table_row(lines[i]))
                i += 1
            if header or rows:
                blocks.append(("table", (header, rows)))
            continue
        ordered_match = MD_ORDERED_LIST_RE.match(line)
        if ordered_match:
            blocks.append(("olist", ordered_match.group(2).strip()))
            i += 1
            continue
        unordered_match = MD_UNORDERED_LIST_RE.match(line)
        if unordered_match:
            blocks.append(("ulist", unordered_match.group(1).strip()))
            i += 1
            continue
        if is_heading_candidate(line):
            blocks.append(("heading", (2, line.strip())))
            i += 1
            continue

        paragraph_lines = [line.strip()]
        i += 1
        while i < len(lines):
            peek = lines[i].strip()
            if not peek:
                break
            if peek.startswith("```"):
                break
            if MD_HEADING_LINE_RE.match(peek):
                break
            if "|" in peek and i + 1 < len(lines) and MD_TABLE_SEPARATOR_RE.match(lines[i + 1]):
                break
            if MD_ORDERED_LIST_RE.match(peek) or MD_UNORDERED_LIST_RE.match(peek):
                break
            if is_heading_candidate(peek):
                break
            paragraph_lines.append(peek)
            i += 1
        blocks.append(("paragraph", " ".join(paragraph_lines)))
    return blocks


def add_heading_paragraph(doc: "Document", text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    run = paragraph.add_run(text)
    run.bold = True
    if Pt is not None:
        if level <= 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
    _apply_run_font(run)


def add_list_paragraph(doc: "Document", text: str, ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    try:
        paragraph = doc.add_paragraph(style=style)
        add_runs_with_style(paragraph, text)
    except KeyError:
        paragraph = doc.add_paragraph(style="Normal")
        prefix = "1. " if ordered else "• "
        add_runs_with_style(paragraph, f"{prefix}{text}")


def add_table(doc: "Document", header: List[str], rows: List[List[str]]) -> None:
    if not header and not rows:
        return
    columns = max(len(header), max((len(row) for row in rows), default=0))
    if columns == 0:
        return
    table = doc.add_table(rows=1 + len(rows), cols=columns)
    table.style = "Table Grid"
    for col_index in range(columns):
        cell_text = header[col_index] if col_index < len(header) else ""
        cell = table.rows[0].cells[col_index]
        cell.text = ""
        add_runs_with_style(cell.paragraphs[0], cell_text)
    for row_index, row in enumerate(rows, start=1):
        for col_index in range(columns):
            cell_text = row[col_index] if col_index < len(row) else ""
            cell = table.rows[row_index].cells[col_index]
            cell.text = ""
            add_runs_with_style(cell.paragraphs[0], cell_text)


def _apply_normal_style(doc: "Document") -> None:
    normal_style = doc.styles["Normal"]
    normal_style.font.bold = False
    normal_style.font.name = "SimSun"
    if qn is not None:
        normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def write_text_doc(text: str, output_path: Path) -> None:
    if Document is None:
        print("❌ 缺少依赖 python-docx，请先安装: pip install python-docx")
        return

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _apply_normal_style(doc)

    paragraphs = normalize_text_to_paragraphs(text or "")
    if paragraphs:
        for paragraph_text in paragraphs:
            paragraph = doc.add_paragraph(style="Normal")
            run = paragraph.add_run(paragraph_text)
            run.bold = False
    else:
        paragraph = doc.add_paragraph(style="Normal")
        run = paragraph.add_run("[空白]")
        run.bold = False

    doc.save(output_path)
    print(f"✅ Word 已生成: {output_path}")


def write_markdown_doc(text: str, output_path: Path) -> None:
    if Document is None:
        print("❌ 缺少依赖 python-docx，请先安装: pip install python-docx")
        return

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _apply_normal_style(doc)

    blocks = parse_markdown_blocks(text or "")
    for block_type, payload in blocks:
        if block_type == "blank":
            doc.add_paragraph("")
            continue
        if block_type == "heading":
            level, heading_text = payload
            add_heading_paragraph(doc, heading_text, level)
            continue
        if block_type == "olist":
            add_list_paragraph(doc, payload, ordered=True)
            continue
        if block_type == "ulist":
            add_list_paragraph(doc, payload, ordered=False)
            continue
        if block_type == "table":
            header, rows = payload
            add_table(doc, header, rows)
            continue
        if block_type == "code":
            paragraph = doc.add_paragraph(style="Normal")
            run = paragraph.add_run(payload)
            run.bold = False
            run.font.name = "SimSun"
            if qn is not None:
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
            continue
        if block_type == "paragraph":
            paragraph = doc.add_paragraph(style="Normal")
            add_runs_with_style(paragraph, payload)

    doc.save(output_path)
    print(f"✅ Word 已生成: {output_path}")


def write_word_doc(text_blocks: List[Tuple[str, str]], output_path: Path) -> None:
    if Document is None:
        print("❌ 缺少依赖 python-docx，请先安装: pip install python-docx")
        return

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _apply_normal_style(doc)

    for index, (title, text) in enumerate(text_blocks, start=1):
        title_paragraph = doc.add_paragraph(style="Normal")
        title_run = title_paragraph.add_run(f"{index}. {title}")
        title_run.bold = False
        if text.strip():
            for paragraph_text in normalize_text_to_paragraphs(text):
                paragraph = doc.add_paragraph(style="Normal")
                run = paragraph.add_run(paragraph_text)
                run.bold = False
        else:
            paragraph = doc.add_paragraph(style="Normal")
            run = paragraph.add_run("[空白]")
            run.bold = False

    doc.save(output_path)
    print(f"✅ Word 已生成: {output_path}")
