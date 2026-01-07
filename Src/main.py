import concurrent.futures
import hashlib
import json
import os
import re
import shutil
import subprocess
import time
from pathlib import Path
from typing import Iterable, List, Optional, Set

from doubao_client import DEFAULT_BASE_URL, build_messages, chat_completion
from pdf_reader import extract_text_from_pdf
from ocr_client import (
    ocr_image_path_to_text,
    resolve_ocr_workers,
    resolve_visual_ocr_config,
)
from word_writer import write_markdown_doc, write_text_doc

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

BASE_DIR = Path(__file__).resolve().parent.parent
INPUT_DIR = BASE_DIR / "Input" / "pdfData"
WORD_INPUT_DIR = BASE_DIR / "Input" / "wordData"
PHOTO_INPUT_DIR = BASE_DIR / "Input" / "photoData"
TEXT_DIR = BASE_DIR / "Output" / "text"
PHOTO_TEXT_DIR = TEXT_DIR / "photo_texts"
OUTPUT_DIR = TEXT_DIR / "combined_documents"
OUTPUT_PDF_TXT = OUTPUT_DIR / "综合文档pdf版本.txt"
OUTPUT_WORD_TXT = OUTPUT_DIR / "综合文档word版本.txt"
OUTPUT_MERGED_TXT = OUTPUT_DIR / "综合合并文档.txt"
OUTPUT_PHOTO_TXT = OUTPUT_DIR / "综合文档图片版本.txt"
OUTPUT_KEYINFO_TXT = OUTPUT_DIR / "关键要素信息提取文档要求.txt"
OUTPUT_EVIDENCE_TXT = OUTPUT_DIR / "质证意见文档要求.txt"
OUTPUT_EVIDENCE_SUGGESTIONS_TXT = OUTPUT_DIR / "补充证据建议清单要求文档.txt"
OUTPUT_COMPLAINT_TXT = OUTPUT_DIR / "起诉状要求文档.txt"
OUTPUT_EVIDENCE_CATALOG_TXT = OUTPUT_DIR / "证据目录.txt"
OUTPUT_DISPUTE_FOCUS_TXT = OUTPUT_DIR / "争议识别要求文档.txt"
OUTPUT_LEGAL_BASIS_TXT = OUTPUT_DIR / "法律依据要求文档.txt"
OUTPUT_RISK_STRATEGY_TXT = OUTPUT_DIR / "风险评估报告和应对建议要求文档.txt"
OUTPUT_LITIGATION_STRATEGY_TXT = OUTPUT_DIR / "诉讼策略报告要求文档.txt"
AI_OUTPUT_DIR = OUTPUT_DIR / "ai_documents"
AI_OUTPUT_INDEX = AI_OUTPUT_DIR / "ai_outputs.json"
SNAPSHOT_FILE = OUTPUT_DIR / "input_snapshot.json"
TITLE_WIDTH = 80
CONTROL_CHARS_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
HYPERLINK_LINE_RE = re.compile(r"^HYPERLINK\\b", re.IGNORECASE)
PAGE_FIELD_RE = re.compile(r"^PAGE/NUMPAGES$", re.IGNORECASE)
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"}


def list_pdf_files(input_dir: Path) -> List[Path]:
    if not input_dir.exists():
        print(f"❌ 输入目录不存在: {input_dir}")
        return []

    pdf_files = [
        path for path in input_dir.iterdir()
        if path.is_file() and path.suffix.lower() == ".pdf"
    ]
    if not pdf_files:
        print(f"❌ 没有找到 PDF 文件，请放入目录: {input_dir}")
    return sorted(pdf_files, key=lambda p: p.name.lower())


def list_word_files(input_dir: Path) -> List[Path]:
    if not input_dir.exists():
        print(f"❌ 输入目录不存在: {input_dir}")
        return []

    word_files = [
        path for path in input_dir.iterdir()
        if path.is_file() and path.suffix.lower() in {".doc", ".docx"}
    ]
    if not word_files:
        print(f"❌ 没有找到 Word 文件（.doc/.docx），请放入目录: {input_dir}")
    return sorted(word_files, key=lambda p: p.name.lower())


def list_photo_files(input_dir: Path) -> List[Path]:
    if not input_dir.exists():
        print(f"❌ 输入目录不存在: {input_dir}")
        return []

    photo_files = [
        path for path in input_dir.iterdir()
        if path.is_file() and path.suffix.lower() in IMAGE_EXTS
    ]
    if not photo_files:
        print(f"❌ 没有找到图片文件，请放入目录: {input_dir}")
    return sorted(photo_files, key=lambda p: p.name.lower())


def iter_page_files(folder: Path) -> List[Path]:
    def page_sort_key(path: Path) -> int:
        name = path.stem
        if name.startswith("page_"):
            suffix = name.split("_", 1)[1]
            if suffix.isdigit():
                return int(suffix)
        return 0

    if not folder.exists():
        return []
    page_files = sorted(folder.glob("page_*.txt"), key=page_sort_key)
    return page_files


def iter_page_texts(folder: Path) -> Iterable[str]:
    for page_file in iter_page_files(folder):
        with open(page_file, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read().strip()
        if text:
            yield text


def read_text_from_folder(folder: Path) -> str:
    return "\n".join(iter_page_texts(folder)).strip()


def iter_word_texts(word_path: Path) -> Iterable[str]:
    def clean_text(value: str) -> str:
        return CONTROL_CHARS_RE.sub("", value).strip()

    suffix = word_path.suffix.lower()
    if suffix == ".docx":
        if DocxDocument is None:
            print("❌ 缺少依赖 python-docx，请先安装: pip install python-docx")
            return

        doc = DocxDocument(word_path)
        for paragraph in doc.paragraphs:
            text = clean_text(paragraph.text)
            if text and not HYPERLINK_LINE_RE.match(text) and not PAGE_FIELD_RE.match(text):
                yield text

        for table in doc.tables:
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = " ".join(
                        clean_text(p.text) for p in cell.paragraphs if clean_text(p.text)
                    ).strip()
                    cells.append(cell_text)
                row_text = clean_text("\t".join(cells))
                if row_text and not HYPERLINK_LINE_RE.match(row_text) and not PAGE_FIELD_RE.match(row_text):
                    yield row_text
        return

    if suffix == ".doc":
        textutil_path = shutil.which("textutil")
        if not textutil_path:
            print("❌ 无法读取 .doc 文件，请先安装 LibreOffice 或转换为 .docx")
            return
        result = subprocess.run(
            [textutil_path, "-convert", "txt", "-stdout", str(word_path)],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="ignore",
        )
        if result.returncode != 0:
            message = result.stderr.strip() or "转换失败"
            print(f"❌ .doc 解析失败: {word_path.name}（{message}）")
            return
        for line in result.stdout.splitlines():
            line = clean_text(line)
            if line and not HYPERLINK_LINE_RE.match(line) and not PAGE_FIELD_RE.match(line):
                yield line
        return

    print(f"❌ 不支持的 Word 格式: {word_path.name}")


def format_title_line(filename: str) -> str:
    title = filename.strip()
    if len(title) >= TITLE_WIDTH:
        return title
    return title.center(TITLE_WIDTH)


def load_env_file(env_path: Path) -> None:
    if not env_path.exists():
        return
    for raw_line in env_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


def get_env_value(name: str) -> Optional[str]:
    value = os.environ.get(name)
    if not value or value.strip() == "":
        return None
    return value.strip()


def get_doubao_settings() -> dict:
    model = get_env_value("DOUBAO_MODEL")
    if not model:
        raise RuntimeError("Missing DOUBAO_MODEL environment variable.")
    return {
        "model": model,
        "base_url": get_env_value("DOUBAO_BASE_URL") or DEFAULT_BASE_URL,
        "reasoning_effort": get_env_value("DOUBAO_REASONING_EFFORT") or "medium",
    }


def resolve_doubao_workers() -> int:
    raw_value = get_env_value("DOUBAO_MAX_WORKERS")
    if raw_value:
        try:
            value = int(raw_value)
            if value > 0:
                return value
        except ValueError:
            pass
    return 2


def resolve_doubao_retry_settings() -> tuple[int, float]:
    retries_raw = get_env_value("DOUBAO_RETRY_TIMES")
    backoff_raw = get_env_value("DOUBAO_RETRY_BACKOFF_SECONDS")
    retries = 2
    backoff = 1.5
    if retries_raw:
        try:
            retries = max(0, int(retries_raw))
        except ValueError:
            pass
    if backoff_raw:
        try:
            backoff = max(0.1, float(backoff_raw))
        except ValueError:
            pass
    return retries, backoff


def format_duration(seconds: float) -> str:
    seconds = max(0.0, seconds)
    if seconds < 60:
        return f"{seconds:.1f}s"
    minutes = seconds / 60
    if minutes < 60:
        return f"{minutes:.1f}m"
    hours = minutes / 60
    return f"{hours:.1f}h"


def hash_text(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def model_signature(settings: dict) -> str:
    return hash_text(f"{settings['model']}|{settings['reasoning_effort']}")


def load_ai_output_index() -> dict:
    if not AI_OUTPUT_INDEX.exists():
        return {}
    try:
        return json.loads(AI_OUTPUT_INDEX.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return {}


def save_ai_output_index(index: dict) -> None:
    AI_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    AI_OUTPUT_INDEX.write_text(
        json.dumps(index, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )


def unique_name(base_name: str, used_names: Set[str], fallback: str = "item") -> str:
    base = base_name.strip() or fallback
    candidate = base
    index = 1
    while candidate.lower() in used_names:
        candidate = f"{base}_{index}"
        index += 1
    used_names.add(candidate.lower())
    return candidate


def resolve_output_folder(pdf_file: Path, used_names: Set[str]) -> Path:
    candidate = unique_name(pdf_file.stem, used_names, fallback="pdf")
    return TEXT_DIR / candidate


def resolve_output_text_file(base_name: str, used_names: Set[str], folder: Path) -> Path:
    candidate = unique_name(base_name, used_names, fallback="text")
    return folder / f"{candidate}.txt"


def write_combined_pdf_txt(pdf_files: List[Path]) -> None:
    if not pdf_files:
        return

    TEXT_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    used_folder_names: Set[str] = {OUTPUT_DIR.name.lower()}

    with open(OUTPUT_PDF_TXT, "w", encoding="utf-8") as output_file:
        for index, pdf_file in enumerate(pdf_files):
            output_folder = resolve_output_folder(pdf_file, used_folder_names)
            extract_text_from_pdf(str(pdf_file), str(output_folder))

            if index > 0:
                output_file.write("\n\n")

            output_file.write(f"{format_title_line(pdf_file.name)}\n")

            first_page = True
            for page_text in iter_page_texts(output_folder):
                if not first_page:
                    output_file.write("\n")
                output_file.write(page_text)
                first_page = False
    print(f"✅ 综合文档 PDF 版本已生成: {OUTPUT_PDF_TXT}")


def write_combined_word_txt(word_files: List[Path]) -> None:
    if not word_files:
        return

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    with open(OUTPUT_WORD_TXT, "w", encoding="utf-8") as output_file:
        for index, word_file in enumerate(word_files):
            if index > 0:
                output_file.write("\n\n")

            output_file.write(f"{format_title_line(word_file.name)}\n")

            first_line = True
            for line in iter_word_texts(word_file):
                if not first_line:
                    output_file.write("\n")
                output_file.write(line)
                first_line = False
    print(f"✅ 综合文档 Word 版本已生成: {OUTPUT_WORD_TXT}")


def ocr_image_to_text(image_path: Path, config: Optional[dict] = None) -> str:
    return ocr_image_path_to_text(image_path, config)


def write_combined_photo_txt(photo_files: List[Path]) -> None:
    if not photo_files:
        return

    config = resolve_visual_ocr_config()
    if not config:
        return

    PHOTO_TEXT_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    used_names: Set[str] = set()
    workers = resolve_ocr_workers()

    results: List[str] = [""] * len(photo_files)
    with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as executor:
        future_map = {
            executor.submit(ocr_image_to_text, photo_file, config): index
            for index, photo_file in enumerate(photo_files)
        }
        for future in concurrent.futures.as_completed(future_map):
            index = future_map[future]
            try:
                results[index] = future.result() or ""
            except Exception as exc:
                print(f"❌ OCR 识别失败: {photo_files[index].name}（{exc}）")
                results[index] = ""

    with open(OUTPUT_PHOTO_TXT, "w", encoding="utf-8") as output_file:
        for index, photo_file in enumerate(photo_files):
            if index > 0:
                output_file.write("\n\n")

            output_file.write(f"{format_title_line(photo_file.name)}\n")
            text = results[index]
            if text:
                output_file.write(text)

            output_text_path = resolve_output_text_file(photo_file.stem, used_names, PHOTO_TEXT_DIR)
            output_text_path.write_text(text, encoding="utf-8")

    print(f"✅ 综合文档 图片版本已生成: {OUTPUT_PHOTO_TXT}")


def write_merged_txt(files: List[Path], output_path: Path) -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    wrote_any = False
    with open(output_path, "w", encoding="utf-8") as output_file:
        for source_path in files:
            if not source_path.exists() or source_path.stat().st_size == 0:
                continue
            if wrote_any:
                output_file.write("\n\n")
            with open(source_path, "r", encoding="utf-8", errors="ignore") as source_file:
                shutil.copyfileobj(source_file, output_file)
            wrote_any = True
    if wrote_any:
        print(f"✅ 综合合并文档已生成: {output_path}")


def write_prompt_doc(merged_path: Path, output_path: Path, prompt: str, label: str) -> None:
    if not merged_path.exists():
        return
    content = merged_path.read_text(encoding="utf-8", errors="ignore").strip()
    if not content:
        return
    output_path.write_text(f"{prompt}\n{content}", encoding="utf-8")
    print(f"✅ {label}已生成: {output_path}")


def convert_docx_to_doc(docx_path: Path, doc_path: Path) -> bool:
    textutil_path = shutil.which("textutil")
    if not textutil_path:
        print(f"⚠️ 未找到 textutil，无法生成 .doc：{doc_path}")
        return False
    result = subprocess.run(
        [textutil_path, "-convert", "doc", "-output", str(doc_path), str(docx_path)],
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="ignore",
    )
    if result.returncode != 0:
        message = result.stderr.strip() or "转换失败"
        print(f"⚠️ .doc 转换失败: {docx_path.name}（{message}）")
        return False
    return True


def generate_ai_documents(ai_docs: List[dict], settings: Optional[dict] = None) -> None:
    if not os.environ.get("ARK_API_KEY"):
        print("⚠️ 缺少 ARK_API_KEY，跳过豆包生成。")
        return
    if settings is None:
        try:
            settings = get_doubao_settings()
        except RuntimeError as exc:
            print(f"⚠️ {exc}，跳过豆包生成。")
            return

    AI_OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    index = load_ai_output_index()
    retries, backoff = resolve_doubao_retry_settings()
    workers = resolve_doubao_workers()
    tasks = []

    signature = model_signature(settings)
    for doc in ai_docs:
        source_path = doc["input"]
        if not source_path.exists():
            print(f"⚠️ 缺少输入文档，跳过生成: {source_path}")
            continue
        content = source_path.read_text(encoding="utf-8", errors="ignore").strip()
        if not content:
            print(f"⚠️ 输入内容为空，跳过生成: {source_path}")
            continue
        output_paths = [
            path
            for path in (
                doc.get("output_docx"),
                doc.get("output_doc"),
                doc.get("output_md"),
            )
            if path is not None
        ]
        if not output_paths:
            continue
        cache_key = f"{source_path}|{signature}"
        content_hash = hash_text(content)
        cached = index.get(cache_key, {})
        output_missing = any(not path.exists() for path in output_paths)
        if not output_missing and cached.get("hash") == content_hash:
            continue
        tasks.append({
            "doc": doc,
            "content": content,
            "hash": content_hash,
            "cache_key": cache_key,
        })

    if not tasks:
        print("ℹ️ 豆包输出已是最新，跳过生成。")
        return
    total_tasks = len(tasks)
    print(f"⏳ AI大模型生成任务数: {total_tasks}，并发: {workers}")

    def call_doubao(text: str) -> str:
        attempt = 0
        while True:
            try:
                messages = build_messages(text)
                return chat_completion(
                    messages,
                    model=settings["model"],
                    base_url=settings["base_url"],
                    reasoning_effort=settings["reasoning_effort"],
                )
            except Exception:
                if attempt >= retries:
                    raise
                time.sleep(backoff * (2 ** attempt))
                attempt += 1

    results = []
    start_time = time.monotonic()
    durations: List[float] = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=workers) as executor:
        future_map = {
            executor.submit(call_doubao, task["content"]): {
                "task": task,
                "start_time": time.monotonic(),
            }
            for task in tasks
        }
        for future in concurrent.futures.as_completed(future_map):
            record = future_map[future]
            task = record["task"]
            task_duration = time.monotonic() - record["start_time"]
            durations.append(task_duration)
            try:
                response = future.result() or ""
            except Exception as exc:
                source_path = task["doc"]["input"]
                print(f"❌ 豆包调用失败: {source_path.name}（{exc}）")
                continue
            completed = len(durations)
            elapsed = time.monotonic() - start_time
            avg_duration = sum(durations) / len(durations)
            remaining = total_tasks - completed
            eta = (remaining / max(1, workers)) * avg_duration
            source_path = task["doc"]["input"]
            print(
                f"✅ 完成 {completed}/{total_tasks}: {source_path.name}，"
                f"已用时 {format_duration(elapsed)}，预计剩余 {format_duration(eta)}"
            )
            results.append({**task, "response": response})

    for result in results:
        doc = result["doc"]
        output_md = doc.get("output_md")
        if output_md is not None:
            output_md.write_text(response, encoding="utf-8")

        output_docx = doc.get("output_docx")
        output_doc = doc.get("output_doc")
        response = result["response"]
        if output_docx is not None:
            write_markdown_doc(response, output_docx)
        if output_doc:
            docx_source = output_docx
            if docx_source is None:
                docx_source = AI_OUTPUT_DIR / f"_tmp_{output_doc.stem}.docx"
                write_markdown_doc(response, docx_source)
            if convert_docx_to_doc(docx_source, output_doc):
                if docx_source.name.startswith("_tmp_"):
                    docx_source.unlink(missing_ok=True)
        index[result["cache_key"]] = {
            "hash": result["hash"],
            "updated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
        }

    save_ai_output_index(index)


def load_snapshot(snapshot_path: Path) -> dict:
    if not snapshot_path.exists():
        return {}
    try:
        return json.loads(snapshot_path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return {}


def save_snapshot(snapshot: dict, snapshot_path: Path) -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    snapshot_path.write_text(
        json.dumps(snapshot, ensure_ascii=True, indent=2),
        encoding="utf-8",
    )


def cleanup_outputs() -> None:
    if OUTPUT_DIR.exists():
        shutil.rmtree(OUTPUT_DIR)
    if not TEXT_DIR.exists():
        return
    for entry in TEXT_DIR.iterdir():
        if entry == OUTPUT_DIR:
            continue
        if entry.is_dir():
            shutil.rmtree(entry)
        elif entry.is_file() and entry.suffix.lower() == ".txt":
            entry.unlink()


def main() -> None:
    start_time = time.monotonic()
    load_env_file(BASE_DIR / ".env")
    pdf_files = list_pdf_files(INPUT_DIR)
    word_files = list_word_files(WORD_INPUT_DIR)
    photo_files = list_photo_files(PHOTO_INPUT_DIR)
    prompt_docs = [
        (OUTPUT_KEYINFO_TXT, "关键要素信息提取文档要求", "PROMPT_KEYINFO"),
        (OUTPUT_EVIDENCE_TXT, "质证意见文档要求", "PROMPT_EVIDENCE"),
        (OUTPUT_EVIDENCE_SUGGESTIONS_TXT, "补充证据建议清单要求文档", "PROMPT_EVIDENCE_SUGGESTIONS"),
        (OUTPUT_COMPLAINT_TXT, "起诉状要求文档", "PROMPT_COMPLAINT_PETITION"),
        (OUTPUT_EVIDENCE_CATALOG_TXT, "证据目录", "PROMPT_EVIDENCE_CATALOG"),
        (OUTPUT_DISPUTE_FOCUS_TXT, "争议识别要求文档", "PROMPT_DISPUTE_FOCUS"),
        (OUTPUT_LEGAL_BASIS_TXT, "法律依据要求文档", "PROMPT_LEGAL_BASIS"),
        (OUTPUT_RISK_STRATEGY_TXT, "风险评估报告和应对建议要求文档", "PROMPT_RISK_STRATEGY"),
        (OUTPUT_LITIGATION_STRATEGY_TXT, "诉讼策略报告要求文档", "PROMPT_LITIGATION_STRATEGY"),
    ]
    ai_docs = [
        {
            "input": OUTPUT_EVIDENCE_SUGGESTIONS_TXT,
            "output_docx": AI_OUTPUT_DIR / "补充证据建议清单文档_AI.docx",
            "output_md": AI_OUTPUT_DIR / "补充证据建议清单文档_AI.md",
        },
        {
            "input": OUTPUT_EVIDENCE_CATALOG_TXT,
            "output_docx": AI_OUTPUT_DIR / "证据目录_AI.docx",
            "output_md": AI_OUTPUT_DIR / "证据目录_AI.md",
        },
        {
            "input": OUTPUT_COMPLAINT_TXT,
            "output_docx": AI_OUTPUT_DIR / "起诉状_AI.docx",
            "output_md": AI_OUTPUT_DIR / "起诉状_AI.md",
        },
        {
            "input": OUTPUT_DISPUTE_FOCUS_TXT,
            "output_docx": AI_OUTPUT_DIR / "争议识别要求文档_AI.docx",
            "output_md": AI_OUTPUT_DIR / "争议识别要求文档_AI.md",
        },
        {
            "input": OUTPUT_LEGAL_BASIS_TXT,
            "output_docx": AI_OUTPUT_DIR / "法律依据要求文档_AI.docx",
            "output_md": AI_OUTPUT_DIR / "法律依据要求文档_AI.md",
        },
        {
            "input": OUTPUT_RISK_STRATEGY_TXT,
            "output_docx": AI_OUTPUT_DIR / "风险评估报告和应对建议要求文档_AI.docx",
            "output_md": AI_OUTPUT_DIR / "风险评估报告和应对建议要求文档_AI.md",
        },
        {
            "input": OUTPUT_LITIGATION_STRATEGY_TXT,
            "output_docx": AI_OUTPUT_DIR / "诉讼策略报告要求文档_AI.docx",
            "output_md": AI_OUTPUT_DIR / "诉讼策略报告要求文档_AI.md",
        },
    ]
    try:
        ai_settings = get_doubao_settings()
    except RuntimeError:
        ai_settings = None
    prompt_values = {env_key: get_env_value(env_key) for _, _, env_key in prompt_docs}
    current_snapshot = {
        "pdf": sorted(file.name for file in pdf_files),
        "word": sorted(file.name for file in word_files),
        "photo": sorted(file.name for file in photo_files),
        "prompts": prompt_values,
    }
    previous_snapshot = load_snapshot(SNAPSHOT_FILE)
    prev_prompts = previous_snapshot.get("prompts") if isinstance(previous_snapshot, dict) else None
    inputs_changed = (
        current_snapshot.get("pdf") != previous_snapshot.get("pdf")
        or current_snapshot.get("word") != previous_snapshot.get("word")
        or current_snapshot.get("photo") != previous_snapshot.get("photo")
    )
    prompts_changed = prompt_values != (prev_prompts or {})
    prompt_docs_missing = any(not output_path.exists() for output_path, _, _ in prompt_docs)
    ai_docs_missing = any(
        not output_path.exists()
        for doc in ai_docs
        for output_path in (
            doc.get("output_docx"),
            doc.get("output_doc"),
            doc.get("output_md"),
        )
        if output_path is not None
    )
    ai_inputs_changed = False
    if ai_settings is not None:
        signature = model_signature(ai_settings)
        ai_index = load_ai_output_index()
        for doc in ai_docs:
            source_path = doc["input"]
            if not source_path.exists():
                continue
            content = source_path.read_text(encoding="utf-8", errors="ignore").strip()
            if not content:
                continue
            cache_key = f"{source_path}|{signature}"
            cached = ai_index.get(cache_key, {})
            if cached.get("hash") != hash_text(content):
                ai_inputs_changed = True
                break
    if (
        not inputs_changed
        and not prompts_changed
        and not prompt_docs_missing
        and not ai_docs_missing
        and not ai_inputs_changed
    ):
        print("ℹ️ 输入文件未变更，跳过重新生成。")
        return

    if inputs_changed or not OUTPUT_MERGED_TXT.exists():
        cleanup_outputs()

        if pdf_files:
            write_combined_pdf_txt(pdf_files)
        if word_files:
            write_combined_word_txt(word_files)
        if photo_files:
            write_combined_photo_txt(photo_files)
        write_merged_txt([OUTPUT_PDF_TXT, OUTPUT_WORD_TXT, OUTPUT_PHOTO_TXT], OUTPUT_MERGED_TXT)

    for output_path, label, env_key in prompt_docs:
        prompt = prompt_values.get(env_key)
        if not prompt:
            print(f"⚠️ 缺少提示语 {env_key}，跳过生成: {output_path}")
            continue
        write_prompt_doc(OUTPUT_MERGED_TXT, output_path, prompt, label)
    if inputs_changed or prompts_changed or prompt_docs_missing or ai_docs_missing or ai_inputs_changed:
        generate_ai_documents(ai_docs, ai_settings)
    save_snapshot(current_snapshot, SNAPSHOT_FILE)
    total_elapsed = time.monotonic() - start_time
    print(f"✅ 全流程完成，总耗时 {format_duration(total_elapsed)}")


if __name__ == "__main__":
    main()
    
